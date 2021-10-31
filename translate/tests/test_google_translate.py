# coding: UTF-8
from __future__ import absolute_import
from __future__ import unicode_literals

# python stdlib imports
# TODO: use of stdlib functions should be minimised
import os
import json
import logging
import mock
import pytest
from py._path.local import LocalPath

# import document libraries
from translate.translate_base import (
    load_workbook,
    Document,
    Presentation,
)

# import constants
from translate.translate_base import (
    CREDS,
    DOCX_STYLE_PROPERTY,
    # LANGUAGE_PROPERTIES,
    # TRANSLATION_RULES,
)



# TODO: move tests around to reflect where the methods being tested now reside

mock_dict = json.load(open(os.path.join(LocalPath(os.path.split(__file__)[0]).join("data"), 'mock_translations.json')))


def mock_session(*args, **kwargs):
    return


def mock_response(*args, **kwargs):
    return [{'language': 'af', 'name': 'Afrikaans'}]


def mock_translation(self, string):
    return {'translatedText': mock_dict[string]}


@pytest.fixture(autouse=True)
def mock_Client(monkeypatch):
    from google.cloud.translate import Client
    monkeypatch.setattr(Client, 'from_service_account_json', staticmethod(mock_session))
    monkeypatch.setattr(Client, 'get_languages', mock_response)


@pytest.fixture
def TRANSLATION_RULES():
    from translate.translate_base import TRANSLATION_RULES
    TRANSLATION_RULES['ignore_styles'].append('Caption')
    return TRANSLATION_RULES


@pytest.fixture
def LANGUAGE_PROPERTIES():
    LANGUAGE_PROPERTIES = {
        'ja': {
            'font': 'Hiragino Sans W3',
            # could also use 'MS Gothic', and for Windows could only find 'Hiragino Sans GB W3'
            'lang': 'Japanese',
            'overwrite': {
                ' / ': '/', # what was this for?
            }
        }
    }
    return LANGUAGE_PROPERTIES


@pytest.fixture
def style_props():
    # from translate.translate_base import DOCX_STYLE_PROPERTY
    return DOCX_STYLE_PROPERTY


@pytest.fixture
def prepare_history(datadir):
    def history(given, exists):
        if not given:
            return None
        else:
            import os
            import shutil
            history_file = os.path.join(datadir, 'history_file.json')
            if os.path.exists(history_file):
                os.remove(history_file)
            if exists:
                history_file = os.path.join(datadir, 'history_file.json')
                existing_history_file = os.path.join(datadir, 'history_copy.json')
                shutil.copy(existing_history_file, history_file)
            return history_file
    return history


@pytest.fixture
def GoogleTranslate(monkeypatch):
    from translate.translate_base import GoogleTranslate
    monkeypatch.setattr(GoogleTranslate, 'request_translation', mock_translation)
    return GoogleTranslate


from docx.shared import RGBColor
lavender = RGBColor(0xff, 0x99, 0xcc)


from docx.shared import Pt


# TODO: refactor tests using mocking to be more efficient and rely less on the mock_translation


class TestGoogleTranslate():

    @pytest.mark.parametrize("target_lang, kwargs, string, expected, stdout",
                             [
                                 ("fr", {'source_lang': 'en'}, 'French', 'français', ''),
                                 ("ja", {'source_lang': 'en'}, 'Japanese', '日本語', ''),
                                 ("ja", {'source_lang': 'en', 'online': False}, 'English', 'ja(English)', ''),
                                 ("ja", {'online': False}, 'Offline', 'ja(Offline)', ''),
                                 ("ja", {'online': False}, '', '', ''),
                                 ("ja", {'online': True}, '3 &lt; 5', '3 <5', ''), # Google drops the space
                                 ("ja", {'online': True, 'show': 10}, 'Red/Yellow/Green/Grey', '赤/黄/緑/グレー',
                                  'Red/Yellow ...\n'),
                             ]
                             )
    def test_translate(self, GoogleTranslate, capsys, target_lang, kwargs, string, expected,
                       stdout, LANGUAGE_PROPERTIES):
        babelfish = GoogleTranslate(CREDS, target_lang, **kwargs)
        result = babelfish.translate(string)
        out, err = capsys.readouterr()
        assert result == expected
        assert out == stdout

    @pytest.mark.parametrize("target_lang, kwargs, string, expected",
                             [
                                 ("fr", {'online': False}, 'middle\nLF', 'fr(middle)\nfr(LF)'),
                                 ("fr", {'online': False}, '\nleading LF', '\nfr(leading LF)'),
                                 ("fr", {'online': False}, 'trailing LF\n', 'fr(trailing LF)\n'),
                             ]
                             )
    def test_multi_line(self, GoogleTranslate, target_lang, kwargs, string, expected):
        babelfish = GoogleTranslate(CREDS, target_lang, **kwargs)
        result = babelfish.multi_line(string)
        assert result == expected

    @pytest.mark.parametrize("target_lang, kwargs, lang_pair, strings, expected_counts",
                             [
                                 ("fr", {'online': False}, 'None-fr',
                                  ('one', 'two', 'three', 'two', 'two', ''),
                                  [0, 2, 3, 1, 3],
                                  ),
                                 ("fr", {'source_lang': 'en'}, 'en-fr',
                                  ('one', 'two', 'three', 'two', 'two', ''),
                                  [3, 2, 0, 1, 3],
                                  ),
                             ]
                             )
    def test_update_stats(self, GoogleTranslate, target_lang, kwargs, lang_pair, strings, expected_counts):
        babelfish = GoogleTranslate(CREDS, target_lang, **kwargs)
        for string in strings:
            babelfish.translate(string)
        babelfish.update_stats()
        stats = babelfish.stats
        assert lang_pair in stats
        result = [value for (key, value) in sorted(stats[lang_pair].items())]
        assert result == expected_counts

    @pytest.mark.parametrize("target_lang, kwargs, hist_given, hist_exists, use_hist, len_hist",
                             [
                                 ("fr", {'source_lang': 'en'}, True, False, True, 0),
                                 ("fr", {'source_lang': 'en'}, True, True, True, 3),
                                 ("fr", {'online': False}, True, True, False, 0),
                                 ("fr", {'source_lang': 'en'}, False, True, False, 0),
                             ]
                             )
    def test_prepare_translation(self, GoogleTranslate, prepare_history,
                                 target_lang, kwargs, hist_given, hist_exists, use_hist, len_hist):
        history_file = prepare_history(hist_given, hist_exists)
        kwargs.update({'history': history_file})
        babelfish = GoogleTranslate(CREDS, target_lang, **kwargs)
        assert babelfish.history == use_hist
        assert babelfish.cloud_requests == 0
        assert len(babelfish.translated) == len_hist
        assert babelfish.history_file == kwargs['history']

    @pytest.mark.parametrize("target_lang, kwargs, hist_given, hist_exists, strings, expected",
                             [
                                 ("fr", {'source_lang': 'en'}, True, True,
                                  ('one', 'two', 'three', 'two', 'two', '', 'four', 'five'),
                                  {'one': 'un', 'two': 'deux', 'three': 'trois', 'four': 'quatre', 'five': 'cinq'},
                                  ),
                                 ("fr", {'source_lang': 'en'}, True, False,
                                  ('one', 'two', 'three', 'two', 'two', '', 'four', 'five'),
                                  {'one': 'un', 'two': 'deux', 'three': 'trois', 'four': 'quatre', 'five': 'cinq'},
                                  ),
                                 ("fr", {'source_lang': 'en'}, False, True,
                                  ('one', 'two', 'three', 'two', 'two', ''),
                                  {'one': 'un', 'two': 'deux', 'three': 'trois'},
                                  ),
                                 ("fr", {'source_lang': 'en'}, False, False,
                                  ('one', 'two', 'three', 'two', 'two', ''),
                                  {'one': 'un', 'two': 'deux', 'three': 'trois'},
                                  ),
                             ]
                             )
    def test_save_history(self, GoogleTranslate, prepare_history,
                                 target_lang, kwargs, hist_given, hist_exists, strings, expected):
        history_file = prepare_history(hist_given, hist_exists)
        kwargs.update({'history': history_file})
        babelfish = GoogleTranslate(CREDS, target_lang, **kwargs)
        for string in strings:
            babelfish.translate(string)
        babelfish.save_history()
        if kwargs['history']:
            assert os.path.exists(babelfish.history_file)
        if babelfish.history:
            f = open(babelfish.history_file)
            history = json.load(f)
            assert {key: value.lower() for (key, value) in history.items()} == expected
            # .lower() because Google's capitalisation seems inconsistent

    @pytest.mark.parametrize("history, fallback, category, error",
                             [
                                 ('some rubbish', 'some rubbish',
                                  pytest.warns, None),
                                 ('/some rubbish_1', 'history_file.json',
                                  pytest.warns, UserWarning),
                                 ('.', 'history_file.json',
                                  pytest.warns, UserWarning),
                             ]
                             )
    def test_invalid_history_arg(self, GoogleTranslate, history, fallback, category, error):
        with category(error):
            babelfish = GoogleTranslate(CREDS, 'ja', history=history)
            babelfish.save_history()
        if error:
            assert babelfish.history_file == os.path.join(os.environ['HOME'], fallback)
        if os.path.exists(babelfish.history_file):
            os.remove(babelfish.history_file)


@pytest.fixture
def TranslateBase():
    from translate.translate_base import TranslateBase
    return TranslateBase


class TestTranslateBase():
    @pytest.mark.parametrize("target_lang, filename, body, extn, target_arg, target",
                             [
                                 ('fr', 'test_document.docx', 'test_document', 'docx',
                                  'test_document_fr.docx', 'test_document_fr.docx'),
                                 ('fr', 'test_document_v0.1.docx', 'test_document_v0.1', 'docx',
                                  None, 'test_document_v0.1_fr.docx'),
                                 ('fr', 'test_document_v0.1.docx', 'test_document_v0.1', 'docx',
                                  'target.docx', 'target.docx'),
                             ]
                             )
    def test_translate_base(self, GoogleTranslate, TranslateBase, datadir, target_lang, filename,
                               target_arg, target, body, extn):
        babelfish = GoogleTranslate(CREDS, target_lang)
        tg = TranslateBase(datadir, filename, babelfish, target=target_arg)
        assert tg.filename == body
        assert tg.ext == extn
        assert isinstance(tg.translator, GoogleTranslate)
        assert tg.target == target

    @pytest.mark.parametrize('target_lang, bf_kwargs, hist_given, '
                             'filename, tg_kwargs, '
                             'new_target, new_target_lang, exp_hist_attr, '
                             'exp_cross_check, exp_len_trans, exp_warning',
                             [
                                 ('fr', {'source_lang': 'en'}, True,
                                  'test_document.docx', {'cross_check': True},
                                  'test_document_fr_en.docx', 'en', True,
                                  True, 0, None),
                                 # the expected results for the next two won't seem to make sense, but the normal
                                 # call to swap_languages is to do it only if self.cross_check
                                 ('fr', {}, True,
                                  'test_document.docx', {'cross_check': True},
                                  'test_document_fr_None.docx', None, True,
                                  False, 0, UserWarning),
                                 ('fr', {}, False,
                                  'test_document.docx', {'cross_check': True},
                                  'test_document_fr_None.docx', None, False,
                                  False, 0, UserWarning),
                             ]
                             )
    def test_swap_languages(self, GoogleTranslate, TranslateBase, datadir, prepare_history,
                            target_lang, bf_kwargs, hist_given, filename, tg_kwargs,
                            new_target, new_target_lang, exp_hist_attr,
                            exp_cross_check, exp_len_trans, exp_warning):
        history_file = prepare_history(hist_given, True)
        bf_kwargs.update({"history": history_file})
        babelfish = GoogleTranslate(CREDS, target_lang, **bf_kwargs)
        with pytest.warns(exp_warning):
            tg = TranslateBase(datadir, filename, babelfish, **tg_kwargs)
            tg.swap_languages()
        assert tg.target == new_target
        assert tg.translator.target_lang == new_target_lang
        assert tg.translator.history == exp_hist_attr
        if history_file:
            assert tg.translator.history_file == "%s_%s.json" % (
                history_file[:-5], new_target_lang,
            )  # also tests tg.add_lang_to_filename()
        assert tg.cross_check == exp_cross_check
        assert len(tg.translator.translated) == exp_len_trans

    @pytest.mark.parametrize('target_lang, bf_kwargs, '
                             'filename, tg_kwargs, '
                             'altb_fname, altb_kwargs, exp_new_name',
                             [
                                 ('fr', {'source_lang': 'en', 'online': False},
                                  'test_document.docx', {},
                                  'test_document.docx', {'lang':'fr'}, 'test_document_fr.docx'),
                                 ('fr', {'source_lang': 'en', 'online': False},
                                  'test_document.docx', {},
                                  'test_document.docx', {}, 'test_document_fr.docx'),
                                 ('fr', {'source_lang': 'en', 'online': False},
                                  'test_document.docx', {},
                                  'test_document', {'lang': 'fr'}, 'test_document_fr.docx'),
                             ]
                             )
    def test_add_lang_to_filename(self, GoogleTranslate, datadir, target_lang, bf_kwargs,
                                  TranslateBase, filename, tg_kwargs,
                                  altb_fname, altb_kwargs, exp_new_name):
        babelfish = GoogleTranslate(CREDS, target_lang, **bf_kwargs)
        tg = TranslateBase(datadir, filename, babelfish, **tg_kwargs)
        new_name = tg.add_lang_to_filename(altb_fname, **altb_kwargs)
        assert new_name == exp_new_name

    @pytest.mark.parametrize('target_lang, bf_kwargs, hist_given, '
                             'filename, tg_kwargs, '
                             'target, new_target, stat_keys',
                             [
                                 ('fr', {'source_lang': 'en', 'online': True}, True,
                                  'test_document.docx', {'cross_check': True},
                                  'test_document_fr.docx', 'test_document_fr_en.docx',
                                  ['en-fr', 'fr-en']),
                             ]
                             )
    def test_execute(self, GoogleTranslate, TranslateDocx, datadir, prepare_history, target_lang, bf_kwargs,
                     filename, tg_kwargs, hist_given,
                     target, new_target, stat_keys):
        history_file = prepare_history(hist_given, False)
        new_hist_file = "%s_%s.json" % (history_file[:-5], bf_kwargs["source_lang"],)
        target = os.path.join(datadir, target)
        new_target = os.path.join(datadir, new_target)
        bf_kwargs.update({"history": history_file})
        filelist = (history_file, new_hist_file, target, new_target)
        for file in filelist:
            if os.path.exists(file):
                os.remove(file)
        babelfish = GoogleTranslate(CREDS, target_lang, **bf_kwargs)
        te = TranslateDocx(datadir, filename, babelfish, **tg_kwargs)
        for file in filelist:
            assert os.path.exists(file)
            if os.path.exists(file):
                os.remove(file)
        assert list(sorted(te.translator.stats.keys())) == stat_keys


@pytest.fixture
def TranslateText():
    from translate.translate_base import TranslateText
    return TranslateText


class TestTranslateText():
    @pytest.mark.parametrize('target_lang, filename, src_lang, target, exp_len',
                             [
                                 ('fr', 'test_textfile.txt', 'en', 'test_textfile_fr.txt', 10),
                             ]
                             )
    def test_translate(self, GoogleTranslate, TranslateText, datadir, target_lang, filename,
                       src_lang, target, exp_len):
        if os.path.exists(os.path.join(datadir, target)):
            os.remove(os.path.join(datadir, target))
        babelfish = GoogleTranslate(CREDS, target_lang, source_lang=src_lang)
        tg = TranslateText(datadir, filename, babelfish)
        assert os.path.exists(os.path.join(datadir, target))
        assert len(babelfish.translated) == exp_len
        if os.path.exists(os.path.join(datadir, target)):
            os.remove(os.path.join(datadir, target))


@pytest.fixture
def TranslateExcel():
    from translate.translate_base import TranslateExcel
    return TranslateExcel


class TestTranslateExcel():

    def test_translate(self, TranslateExcel, GoogleTranslate, datadir):
        target_language = 'ja'
        babelfish = GoogleTranslate(CREDS, target_language, source_lang='en', online=False)
        filepath = datadir
        filename = 'test_spreadsheet.xlsx'
        TranslateExcel(filepath, filename, babelfish)

        wb = load_workbook(os.path.join(filepath, 'test_spreadsheet_ja.xlsx'))
        assert wb.sheetnames == ['Sheet1', 'Sheet2']
        assert wb['Sheet1']['A1'].value == 'ja(Heading)'
        assert wb['Sheet2']['A1'].value == 'ja(Heading in Sheet2)'
        assert len(babelfish.translated) == 6
        if os.path.exists(os.path.join(datadir, 'test_spreadsheet_ja.xlsx')):
            os.remove(os.path.join(datadir, 'test_spreadsheet_ja.xlsx'))


@pytest.fixture
def TranslateDocx():
    from translate.translate_base import TranslateDocx
    return TranslateDocx


class TestTranslateDocx():

    def test_translate_and_set_language(self, GoogleTranslate, TranslateDocx, datadir,
                                        TRANSLATION_RULES, LANGUAGE_PROPERTIES):
        target_language = 'ja'
        babelfish = GoogleTranslate(CREDS, target_language, source_lang='en', online=False)
        filepath = datadir
        filename = 'test_document.docx'
        TranslateDocx(filepath, filename, babelfish)

        d = Document(os.path.join(filepath, 'test_document_ja.docx'))
        assert len(d.paragraphs) == 42
        assert len(d.paragraphs[24].runs) == 3
        assert len(d.paragraphs[27].runs) == 4
        assert d.paragraphs[34].text == "ja(Code block)"
        assert d.paragraphs[35].text == "def list_languages():"
        assert d.core_properties.language == 'Japanese'
        assert d.styles["Normal"].font.name == 'Hiragino Sans W3'
        if os.path.exists(os.path.join(filepath, 'test_document_ja.docx')):
            os.remove(os.path.join(filepath, 'test_document_ja.docx'))

    @pytest.mark.parametrize("text, style, exp_text",
                             [
                                 ('This style should translate', 'Normal', 'ja(This style should translate)'),
                                 ('This style should not', 'Caption', 'This style should not'),
                             ])
    def test_translate_paragraphs(self, GoogleTranslate, TranslateDocx, datadir,
                                  text, style, exp_text, TRANSLATION_RULES):
        target_language = 'ja'
        babelfish = GoogleTranslate(CREDS, target_language, source_lang='en', online=False)
        filepath = datadir
        filename = 'paragraph_test.docx'
        d = Document()
        d.add_paragraph(text, style=style)
        d.save(os.path.join(filepath, filename))
        td = TranslateDocx(filepath, filename, babelfish, condense=True)
        assert len(td.document.paragraphs) == 1
        assert td.document.paragraphs[0].style.name == style
        assert len(td.document.paragraphs[0].runs) == 1
        assert td.document.paragraphs[0].runs[0].text == exp_text
        for file in ('paragraph_test.docx', 'paragraph_test_ja.docx'):
            if os.path.exists(os.path.join(filepath, file)):
                os.remove(os.path.join(filepath, file))

    # These two tests relate to methods in TranslateBase
    @pytest.mark.parametrize("run_list, condensed_run_list, run_count",
                             [
                                 ((['There are', False], [' five', False], [' runs ', False],
                                   ['in the', False], [' paragraph', False]),
                                  ['ja(There are five runs in the paragraph)', '', '', '', ''], 5),
                                 ((['There are', True], [' five', False], [' runs ', False],
                                   ['in the', False], [' paragraph', False]),
                                  ['ja(There are)', 'ja( five runs in the paragraph)', '', '', ''], 5),
                                 ((['There are', False], [' five', True], [' runs ', False],
                                   ['in the', False], [' paragraph', False]),
                                  ['ja(There are)', 'ja( five)', 'ja( runs in the paragraph)', '', ''], 5),
                                 ((['There are', False], [' five', True], [' runs ', True],
                                   ['in the', False], [' paragraph', False]),
                                  ['ja(There are)', 'ja( five runs )', '', 'ja(in the paragraph)', ''], 5),
                                 ((['', False], ['There are five', False], [' runs ', False],
                                   ['in the', False], [' paragraph', False]),
                                  ['', 'ja(There are five runs in the paragraph)', '', '', ''], 5),
                                 ((['There are', False], [' five', False], [' runs ', False],
                                   ['in \nthe', False], [' paragraph', False]),
                                  ['ja(There are five runs in )\nja(the paragraph)', '', '', '', ''], 5),
                             ]
                             )
    def test_condense_runs(self, TranslateDocx, GoogleTranslate, datadir,
                           run_list, condensed_run_list, run_count):
        target_language = 'ja'
        babelfish = GoogleTranslate(CREDS, target_language, source_lang='en', online=False)
        filepath = datadir
        filename = 'condense.docx'
        d = Document()
        p = d.add_paragraph()
        for s, emp in run_list:
            if emp:
                p.add_run(s).bold = True
            else:
                p.add_run(s)
        d.save(os.path.join(filepath, filename))

        TranslateDocx(filepath, filename, babelfish, condense=True)
        d = Document(os.path.join(filepath, 'condense_ja.docx'))
        assert len(d.paragraphs) == 1
        assert len(d.paragraphs[0].runs) == run_count  # ideally this would reduce rather than having empty runs
        for i in range(run_count):
            assert d.paragraphs[0].runs[i].text == condensed_run_list[i]
        for file in ('condense.docx', 'condense_ja.docx'):
            if os.path.exists(os.path.join(filepath, file)):
                os.remove(os.path.join(filepath, file))

    @pytest.mark.parametrize('para_mod, run_list, same',
                             [
                                 # first group applies the style directly to the run
                                 (None,
                                  [['Plain, ', None, None], ['Same', None, None]],
                                  True),
                                 (['style.font.bold', True],
                                  [['Bold_from_para ', None, None], ['Same', None, None]],
                                  True),
                                 (None,
                                  [['Plain ', None, None], ['Bold', 'bold', True]],
                                  False),
                                 (['style.font.bold', True],
                                  [['Bold_from_para ', None, None], ['Bold', 'bold', True]],
                                  True),
                                 (None,
                                  [['Bold', 'bold', True], ['Same', 'bold', True]],
                                  True),
                                 (['style.font.bold', True],
                                  [['Bold', 'bold', True], ['Same', 'bold', True]],
                                  True),
                                 (None,
                                  [['Plain ', None, None], ['Italic', 'italic', True]],
                                  False),
                                 (['style.font.italic', True],
                                  [['Italic_from_para ', None, None], ['Italic', 'italic', True]],
                                  True),
                                 (None,
                                  [['Underline', 'underline', True], ['Plain ', None, None]],
                                  False),
                                 (['style.font.underline', True],
                                  [['Underline', 'underline', True], ['Underline_from_para ', None, None]],
                                  True),
                                 # second group applies the style using the run.font attribute
                                 (None,
                                  [['Plain ', None, None], ['Bold', 'font.bold', True]],
                                  False),
                                 (['style.font.bold', True],
                                  [['Bold_from_para ', None, None], ['Bold', 'font.bold', True]],
                                  True),
                                 (None,
                                  [['Plain ', 'font.bold', False], ['Bold', 'font.bold', True]],
                                  False),
                                 (['style.font.bold', True],
                                  [['Plain ', 'font.bold', False], ['Bold', 'font.bold', True]],
                                  False),
                                 (None,
                                  [['Plain ', None, None], [' Small', 'font.size', Pt(6)]],
                                  False),
                                 (['style.font.size', Pt(6)],
                                  [['Small_from_para ', None, None], [' Small', 'font.size', Pt(6)]],
                                  True),
                                 (None,
                                  [['Large ', 'font.size', Pt(24)], [' Small', 'font.size', Pt(6)]],
                                  False),
                                 (None,
                                  [['Plain ', None, None], ['Caps', 'font.all_caps', True]],
                                  False),
                                 (None,
                                  [['Plain ', None, None], [' Courier', 'font.name', 'Courier']],
                                  False),
                                 # Next one paragraph font name is overwritten in the translation process
                                 (['style.font.name', 'Courier'],
                                  [['Hiragino Sans W3', None, None], [' Courier', 'font.name', 'Courier']],
                                  False),
                                 # next one follows the font.color branch
                                 (None,
                                  [['Plain ', None, None], ['White', 'font.color.rgb', lavender]],
                                  False),
                                 (['style.font.color.rgb', lavender],
                                  [['Lavender_from_para ', None, None], ['White', 'font.color.rgb', lavender]],
                                  True),
                                 # this is the last one in the DOCX_STYLE_PROPERTY to test the full tree
                                 (None,
                                  [['Plain ', None, None], [' Invisible', 'font.spec_vanish', True]],
                                  False),
                             ])
    def test_same_style_runs(self, TranslateDocx, GoogleTranslate, datadir, style_props,
                             para_mod, run_list, same):
        target_language = 'ja'
        babelfish = GoogleTranslate(CREDS, target_language, source_lang='en', online=False)
        filepath = datadir
        filename = 'style_test.docx'
        d = Document()
        p = d.add_paragraph(style='Normal')
        if para_mod:
            self.setattr_drill(p, para_mod[0], para_mod[1])
        for s, attr, val in run_list:
            if attr:
                run = p.add_run(s)
                self.setattr_drill(run, attr, val)
            else:
                p.add_run(s)
        d.save(os.path.join(filepath, filename))

        td = TranslateDocx(filepath, filename, babelfish)
        args = [td.document.paragraphs[0].runs[i] for i in range(2)]
        args.extend([td.document.paragraphs[0].style, style_props])
        result = td.same_style_runs(*args)
        assert result == same

        for file in (filename, 'style_test_ja.docx'):
            if os.path.exists(os.path.join(filepath, file)):
                os.remove(os.path.join(filepath, file))

    def setattr_drill(self, o, attr_str, value):
        if '.' in attr_str:
            attr_list = attr_str.split('.')
            self.setattr_drill(getattr(o, attr_list[0]), '.'.join(attr_list[1:]), value)
        else:
            setattr(o, attr_str, value)


@pytest.fixture
def TranslatePptx():
    from translate.translate_base import TranslatePptx
    return TranslatePptx


class TestTranslatePptx():
    def test_translate_and_set_language(self, GoogleTranslate, TranslatePptx, datadir,
                                        TRANSLATION_RULES, LANGUAGE_PROPERTIES):
        target_language = 'ja'
        babelfish = GoogleTranslate(CREDS, target_language, source_lang='en', show=0, online=False)
        filepath = datadir
        filename = 'test_presentation.pptx'
        TranslatePptx(filepath, filename, babelfish, condense=True)

        prs = Presentation(os.path.join(filepath, 'test_presentation_ja.pptx'))
        assert len(prs.slides) == 7
        assert len(prs.slides[0].shapes) == 2
        assert len(prs.slides[1].shapes) == 5
        assert len(prs.slides[2].shapes) == 2
        assert len(prs.slides[3].shapes) == 2
        assert len(prs.slides[4].shapes) == 2
        assert len(prs.slides[5].shapes) == 2
        assert len(prs.slides[6].shapes) == 2
        assert len(prs.slides[1].shapes[1].text_frame.paragraphs) == 5
        test_pars = prs.slides[1].shapes[1].text_frame.paragraphs
        assert test_pars[0].runs[0].text == 'ja(Bullet 1: no style changes)'
        assert test_pars[1].runs[0].text == 'ja(Bullet 2: style )'
        assert test_pars[2].runs[1].text == 'ja(bold)'
        assert test_pars[3].runs[2].text == 'ja( in the middle)'
        assert test_pars[4].runs[0].text == 'ja(Bullet 5: has a )'
        assert prs.core_properties.language == 'Japanese'
        test_cells = prs.slides[5].shapes[1].table.rows[1].cells
        assert test_cells[0].text_frame.paragraphs[0].runs[0].text == 'ja(First row text in first column)'
        assert test_cells[1].text_frame.paragraphs[0].runs[0].text == 'ja(First row text )'
        assert test_cells[2].text_frame.paragraphs[0].runs[0].text == 'ja(First row text )'
        assert test_cells[3].text_frame.paragraphs[0].runs[0].text == 'ja(First row text with )'
        if os.path.exists(os.path.join(filepath, 'test_presentation_ja.pptx')):
            os.remove(os.path.join(filepath, 'test_presentation_ja.pptx'))

    @pytest.mark.parametrize('expected',
                             [
                                 ([[False, True],
                                   [False, True],
                                   [True, True],
                                   [False, True],  # there is no flag for line feed at end
                                   None]),  # this is only a single run
                             ])
    def test_break_runs(selfself, TranslatePptx, GoogleTranslate, datadir, style_props, expected):
        target_language = 'ja'
        babelfish = GoogleTranslate(CREDS, target_language, source_lang='en', online=False)
        filepath = datadir
        filename = 'test_presentation.pptx'
        tp = TranslatePptx(filepath, filename, babelfish, condense=True)
        test_pars = tp.prs.slides[4].shapes[1].text_frame.paragraphs
        for i, paragraph in enumerate(test_pars):
            break_list = tp.break_runs(paragraph)
            assert break_list == expected[i]
        if os.path.exists(os.path.join(filepath, 'test_presentation_ja.pptx')):
            os.remove(os.path.join(filepath, 'test_presentation_ja.pptx'))

    @pytest.mark.parametrize('break_list, expected',
                             [
                                 ([False, True, False],
                                  [False, True]),
                                 ([False, False, True, False],
                                  [False, False, True]),
                                 ([False, True, False, True, False],
                                  [False, True, True]),
                                 ([False, True, False, False],
                                  [False, True, False]),
                                 ([False, True, True, False, False],
                                  [False, True, False]),
                             ]
                             )
    def test_break_at_run(self, GoogleTranslate, TranslatePptx, datadir, break_list, expected):
        target_language = 'ja'
        babelfish = GoogleTranslate(CREDS, target_language, source_lang='en', show=0, online=False)
        filepath = datadir
        filename = 'test_presentation.pptx'
        tp = TranslatePptx(filepath, filename, babelfish)
        assert list(tp.break_at_run(break_list)) == expected
        if os.path.exists(os.path.join(filepath, 'test_presentation_ja.pptx')):
            os.remove(os.path.join(filepath, 'test_presentation_ja.pptx'))

@pytest.fixture()
def parse_args():
    from translate.__main__ import parse_args
    return parse_args


class TestParser():
    @pytest.mark.parametrize('os_call, expected, category, error',
                             [
                                 ('translate translate/tests/data/test_textfile.txt fr',
                                  {'file': 'translate/tests/data/test_textfile.txt',
                                   'target_lang': 'fr', 'source_lang': None, 'online': True,
                                   'show': 0, 'history': None, 'target': None, 'condense': False,
                                   'cross_check': False, 'verbose': 0, 'quiet': 0},
                                  pytest.warns, None),
                                 ('translate -vvq translate/tests/data/test_textfile.txt fr',
                                  {'verbose': 2, 'quiet': 1},
                                  pytest.warns, None),
                                 ('translate -s en -r textfile_fr.json textfile.txt fr',
                                  {'source_lang': 'en', 'history': 'textfile_fr.json'},
                                  pytest.warns, None),
                                 ('translate -cx textfile.txt fr',
                                  {'condense': True, 'cross_check': True},
                                  pytest.warns, None),
                                 ('translate -n -p 20 textfile.txt fr',
                                  {'online': False, 'show': 20},
                                  pytest.warns, None),
                                 ('translate -d textfile_fr.txt  textfile.txt fr -l',
                                  {'target': 'textfile_fr.txt', 'list_langs': True},
                                  pytest.warns, None),
                                 ('translate -j textfile.txt fr',
                                  {'file': 'translate/tests/data/test_textfile.txt'},
                                  pytest.raises, SystemExit),
                                 ('translate textfile.txt',
                                  {'file': 'textfile.txt'},
                                  pytest.raises, SystemExit),
                                 ('translate textfile.txt fr -v',
                                  {'file': 'textfile.txt', 'verbose': 1},
                                  pytest.warns, None),
                                 ('translate -s en textfile.txt fr -v',
                                  {'file': 'textfile.txt', 'verbose': 1, 'source_lang': 'en'},
                                  pytest.warns, None),
                             ])
    def test_parse_args(self, parse_args, os_call, expected, category, error):
        arg_list = os_call.split()[1:]
        with category(error):
            args = parse_args(arg_list, '')
            for arg in expected:
                assert getattr(args, arg) == expected[arg]


@pytest.fixture
def ListLanguages(monkeypatch):
    from translate.__main__ import ListLanguages
    monkeypatch.setattr(ListLanguages, 'request_languages', mock_response)
    return ListLanguages


class TestListLanguages():
    def test_list_languages(self, ListLanguages, capsys):
        ListLanguages()
        out, err = capsys.readouterr()
        assert out == 'Afrikaans (af)\n'


@pytest.fixture
def correct_rel_path():
    def prepend_rel_path(arg_list):
        rel_path = os.path.relpath(os.path.split(__file__)[0], os.getcwd())
        rel_args = [
            os.path.join(rel_path, arg) if _ == 0 else arg
            for _, arg in (enumerate(arg_list))
        ]
        return rel_args
    return prepend_rel_path

@pytest.mark.parametrize('arg_list, log_level, category, error, bf_kwargs, filename',
                         [
                             (['data/test_textfile.txt', 'fr', '-v'], logging.INFO, pytest.warns, None,
                              dict(history=None, online=True, show=0, source_lang=None),
                              'test_textfile.txt'),
                             (['textfile.txt', 'fr'], logging.WARNING, pytest.warns, UserWarning,
                              dict(history=None, online=True, show=0, source_lang=None),
                              None),
                         ]
                         )
def test_main(GoogleTranslate, TranslateBase, TranslateText, datadir, correct_rel_path,
              arg_list, log_level, category, error, bf_kwargs, filename):
    arg_list = correct_rel_path(arg_list)
    log = logging.getLogger('translate.__main__')
    del log.manager.loggerDict['translate.__main__']
    with mock.patch.object(GoogleTranslate, "__init__", return_value=None) as gti:
        with mock.patch.object(TranslateText, "__init__", return_value=None) as tti:
            with mock.patch.object(TranslateBase, "execute", return_value=None) as tte:
                from translate.__main__ import main
                with category(error):
                    main(arg_list=arg_list)
    if not error:
        gti.assert_called_once_with(CREDS, 'fr', **bf_kwargs)
        assert tti.call_args[0][0] == datadir
        assert tti.call_args[0][1] == filename
        log = logging.getLogger('translate.__main__')
        assert log.getEffectiveLevel() == log_level

    filename = ('_%s.' % arg_list[1]).join(arg_list[0].rsplit('.', 1))
    if os.path.exists(filename):
        os.remove(filename)

def test_init():
    from translate import __main__ as module
    with mock.patch.object(module, "main", return_value=42):
        with mock.patch.object(module, "__name__", "__main__"):
            with mock.patch.object(module.sys, 'exit') as mock_exit:
                module.init()
                assert mock_exit.call_args[0][0] == 42

